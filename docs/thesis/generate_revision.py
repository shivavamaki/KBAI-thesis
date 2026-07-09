"""
Generates Thesis_Proposal_Revised.docx implementing all 8 reviewer instructions.
Run with: C:\Python312\python.exe generate_revision.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins (A4, 2.54 cm all sides) ─────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
for attr in ('left_margin','right_margin','top_margin','bottom_margin'):
    setattr(section, attr, Cm(2.54))

# ── Helper styles ─────────────────────────────────────────────────────────────
normal = doc.styles['Normal']
normal.font.name = 'TH Sarabun New'
normal.font.size = Pt(16)

def h1(text, color=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def h2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def h3(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    return p

def body(text, indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run(text)
    run.font.size = Pt(16)
    return p

def note(text):
    p = doc.add_paragraph()
    run = p.add_run(f'[แก้ไข] {text}')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    run.italic = True
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.25 + level * 0.63)
    run = p.add_run(text)
    run.font.size = Pt(16)
    return p

def divider():
    p = doc.add_paragraph('─' * 70)
    p.runs[0].font.size = Pt(10)
    p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ═══════════════════════════════════════════════════════════════════════════════
# COVER
# ═══════════════════════════════════════════════════════════════════════════════
h1('เอกสารประกอบการเสนอหัวข้อโครงงานปริญญานิพนธ์')
h1('สาขาภาควิชาชีวการแพทย์และสารสนเทศศาสตร์ทางสุขภาพ คณะเภสัชศาสตร์')
h1('มหาวิทยาลัยศิลปากร')
h1('ภาคการศึกษาปลาย ปีการศึกษา 2567')
doc.add_paragraph()

# ── Section 1: Thesis code ────────────────────────────────────────────────────
h3('1.  รหัสโครงงานปริญญานิพนธ์')
body('–')

# ── Section 2: Title (REVISED per instructions 1 & 4) ────────────────────────
doc.add_paragraph()
h3('2.  ชื่อหัวข้อโครงงานปริญญานิพนธ์')
note('ข้อ 4 – ปรับชื่อเป็น "การจำแนกประเภทความผิดพลาดของใบสั่งยา" (เพิ่ม "ความผิดพลาด")')
body(
    'การจำแนกประเภทความผิดพลาดของใบสั่งยาแบบหลายป้ายกำกับ'
    'ด้วย Knowledge Base Generative AI',
    indent=False
)
body(
    'Prescription Error Multi-label Classification '
    'with Knowledge Base Generative AI',
    indent=False
)

# ── Section 3: Researcher ─────────────────────────────────────────────────────
doc.add_paragraph()
h3('3.  ผู้รับผิดชอบโครงงานฯ')
body('ภก.ศิวาวุธ มงคล          รหัสนักศึกษา 660820006', indent=False)

# ── Section 4: Advisor ───────────────────────────────────────────────────────
doc.add_paragraph()
h3('4.  อาจารย์ที่ปรึกษาโครงงาน')
body('รองศาสตราจารย์ ดร. ภก.วีรยุทธ์ เลิศนที', indent=False)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# REVISION SUMMARY
# ═══════════════════════════════════════════════════════════════════════════════
h1('สรุปการแก้ไขตามข้อเสนอแนะของคณะกรรมการ (8 ข้อ)', color=(0x1F, 0x49, 0x7D))
doc.add_paragraph()

revisions = [
    ('ข้อ 1', 'ปรับรูปแบบการเขียนโครงร่างฯ และเพิ่มรายละเอียดกระบวนการวิจัย'),
    ('ข้อ 2', 'ขยายความกระบวนการสร้างระบบ Knowledge Base Generative AI'),
    ('ข้อ 3', 'ชี้แจงการนำเข้าข้อมูล HIS และขอบเขต Prescription Error'),
    ('ข้อ 4', 'ปรับชื่อหัวข้อวิทยานิพนธ์'),
    ('ข้อ 5', 'เพิ่มหลักเกณฑ์การคัดเลือก AI ที่ใช้ทดสอบ'),
    ('ข้อ 6', 'ระบุวิธีการสุ่มตัวอย่างใบสั่งยา'),
    ('ข้อ 7', 'เพิ่มการคำนวณ Precision/F1-Score เปรียบเทียบกับเภสัชกรและ baseline'),
    ('ข้อ 8', 'ลบ Reference ที่ซ้ำกัน (Mesko & Topol 2023b)'),
]
for code, desc in revisions:
    bullet(f'{code}: {desc}')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER 1 – Introduction (revised sections only)
# ═══════════════════════════════════════════════════════════════════════════════
h1('บทที่ 1  บทนำ  (ส่วนที่แก้ไข)')
divider()

# ── Objective (revised per instruction 3 & 4) ────────────────────────────────
h2('วัตถุประสงค์ของการวิจัย  [แก้ไขตามข้อ 3 & 4]')
note('ข้อ 3 – เน้นว่าการวิจัยเน้น Prescription Error เท่านั้น ไม่รวม Medication Error ทั้งหมด')
note('ข้อ 4 – ปรับคำ "จำแนกประเภทใบสั่งยา" เป็น "จำแนกประเภทความผิดพลาดของใบสั่งยา"')

body(
    '1. เพื่อพัฒนาระบบ Knowledge Base Generative AI '
    'สำหรับการจำแนกประเภทความผิดพลาดของใบสั่งยา (Prescription Error Classification) '
    'ในลักษณะแบบหลายป้ายกำกับ (Multi-label Classification) '
    'โดยเฉพาะความผิดพลาดที่เกิดจากกระบวนการสั่งยา ซึ่งไม่รวมความผิดพลาดอื่น ๆ '
    'ในวงจรการใช้ยาทั้งหมด (Medication Error)'
)
body(
    '2. เพื่อประเมินประสิทธิภาพของระบบ Knowledge Base Generative AI '
    'ในการตรวจจับและจำแนกประเภทความผิดพลาดของใบสั่งยาหลายประเภท '
    'ที่อาจเกิดร่วมกันในใบสั่งยาเดียว เปรียบเทียบกับการประเมินโดยเภสัชกร '
    'และเปรียบเทียบกับระบบ AI แบบไม่มีฐานความรู้ (without Knowledge Base)'
)
body(
    '3. เพื่อศึกษาความสอดคล้องระหว่างผลการจำแนกของระบบ AI กับการประเมินของเภสัชกร '
    'วิเคราะห์กรณีที่ผลไม่เหมือนกัน (Discordant Cases) '
    'และวิเคราะห์แนวทางในการปรับปรุงความแม่นยำของระบบ'
)

doc.add_paragraph()

# ── Scope (revised per instruction 3) ─────────────────────────────────────────
h2('ขอบเขตของการวิจัย  [แก้ไขตามข้อ 3]')
note('ข้อ 3 – ระบุชัดเจนว่าขอบเขตคือ Prescription Error ไม่ใช่ Medication Error ทั้งหมด')

body(
    'การวิจัยนี้จะดำเนินการเป็นระยะเวลา 12 เดือน ตั้งแต่วันที่ 1 มีนาคม พ.ศ. 2568 '
    'ถึง 28 กุมภาพันธ์ พ.ศ. 2569 โดยทำการเก็บข้อมูลย้อนหลัง (Retrospective) '
    'เกี่ยวกับความผิดพลาดในกระบวนการสั่งจ่ายยา (Prescription Errors) '
    'ซึ่งหมายถึงความผิดพลาดที่เกิดขึ้นในขั้นตอนการสั่งยาโดยแพทย์เท่านั้น '
    'และไม่รวมความผิดพลาดในขั้นตอนอื่น ๆ ของวงจรการใช้ยา เช่น '
    'การเตรียมยา การจ่ายยา หรือการบริหารยา'
)
body(
    'ข้อมูลถูกดึงจากระบบสารสนเทศโรงพยาบาล (Hospital Information System; HIS) '
    'ของโรงพยาบาลกรุงเทพขอนแก่น ซึ่งจัดเก็บในรูปแบบเวชระเบียนอิเล็กทรอนิกส์ (EHR) '
    'โดยใช้ข้อมูลใบสั่งยาทุกรายการของโรงพยาบาล และประเมินผลลัพธ์โดยใช้ข้อมูล '
    'จากระบบเวชระเบียนอิเล็กทรอนิกส์ (EHR) เพื่อตรวจสอบประสิทธิภาพของ AI '
    'ในการลดข้อผิดพลาดในการสั่งจ่ายยาและเพิ่มความปลอดภัยให้กับผู้ป่วย'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER 3 – Methodology (main revision area)
# ═══════════════════════════════════════════════════════════════════════════════
h1('บทที่ 3  วิธีดำเนินการวิจัย  (ส่วนที่แก้ไข / เพิ่มเติม)')
divider()

# ── Sampling (instruction 6) ──────────────────────────────────────────────────
h2('กลุ่มตัวอย่าง  [แก้ไขตามข้อ 6]')
note('ข้อ 6 – ระบุวิธีการสุ่มตัวอย่าง สถิติที่ใช้คำนวณขนาดกลุ่มตัวอย่าง')

h3('การกำหนดขนาดกลุ่มตัวอย่าง')
body(
    'ขนาดกลุ่มตัวอย่างคำนวณโดยใช้สูตรการประมาณค่าสัดส่วนประชากร (Proportion Estimation) '
    'อ้างอิงจากอัตราความชุกของ Prescription Error ที่ร้อยละ 20 ตามงานวิจัยของ '
    'Corny และคณะ (2020) ณ ระดับความเชื่อมั่น 95% (Z = 1.96) '
    'และค่าความคลาดเคลื่อนที่ยอมรับได้ (margin of error, d) ร้อยละ 5 '
    'ซึ่งได้ขนาดกลุ่มตัวอย่างขั้นต่ำ n ≥ 246 ใบสั่งยา '
    'และเพิ่มอัตราการสูญหายของข้อมูลร้อยละ 10 รวมเป็นไม่น้อยกว่า 270 ใบสั่งยา'
)
body(
    'สูตรคำนวณขนาดกลุ่มตัวอย่าง: n = Z²·p(1−p) / d² '
    'โดย p = 0.20, Z = 1.96, d = 0.05 → n = (1.96)²·(0.20)(0.80) / (0.05)² ≈ 246'
)

h3('วิธีการสุ่มตัวอย่าง')
body(
    'ใช้วิธีการสุ่มตัวอย่างแบบมีระบบ (Systematic Random Sampling) '
    'โดยเรียงลำดับใบสั่งยาทุกใบในช่วงเวลาที่กำหนด (1 มีนาคม 2568 – 28 กุมภาพันธ์ 2569) '
    'ตามลำดับวันที่สั่งยา และคัดเลือกทุก k ใบ โดย k = N/n '
    '(N = จำนวนใบสั่งยาทั้งหมดในช่วงเวลา, n = ขนาดกลุ่มตัวอย่างที่ต้องการ) '
    'เริ่มจากตำแหน่งสุ่ม (random start) ระหว่าง 1 ถึง k'
)
body(
    'กลุ่มตัวอย่างในการศึกษาคือใบสั่งยาทั้งหมดที่ได้รับการสั่งจ่ายภายในโรงพยาบาล '
    'กรุงเทพขอนแก่น ในช่วงระยะเวลาที่กำหนด โดยมีเกณฑ์การคัดเลือกดังต่อไปนี้'
)

h3('เกณฑ์การคัดเข้า')
bullet('เป็นใบสั่งยาที่มีการสั่งจ่ายในช่วงวันที่ 1 มีนาคม พ.ศ. 2568 ถึง 28 กุมภาพันธ์ พ.ศ. 2569')
bullet('เป็นใบสั่งยาที่มีรายการยาตั้งแต่ 5 รายการขึ้นไป (Polypharmacy)')

h3('เกณฑ์การคัดออก')
bullet('เป็นใบสั่งยาที่เกิดจากข้อผิดพลาดของระบบสารสนเทศ')
bullet('เป็นใบสั่งยาที่มีข้อมูลไม่ครบถ้วน เช่น ขาดข้อมูลสำคัญของผู้ป่วย หรือไม่มีข้อมูลผลลัพธ์การวิเคราะห์จากระบบ AI')

doc.add_paragraph()

# ── AI Selection (instruction 5) ──────────────────────────────────────────────
h2('การคัดเลือก AI ที่มาใช้ทดสอบ  [เพิ่มใหม่ตามข้อ 5]')
note('ข้อ 5 – ระบุ AI models ที่จะคัดเลือก เหตุผลในการเลือก และวิธีการเปรียบเทียบ')

body(
    'ในการวิจัยนี้ จะทำการทดสอบและเปรียบเทียบระบบ Knowledge Base Generative AI '
    'โดยใช้โมเดลภาษาขนาดใหญ่ (Large Language Models; LLMs) จาก OpenAI '
    'เพื่อประเมินประสิทธิภาพในการจำแนกประเภทความผิดพลาดของใบสั่งยา'
)

h3('โมเดล AI ที่คัดเลือก')
body('โมเดลที่จะทำการทดสอบและเปรียบเทียบ ได้แก่')
bullet('GPT-4o (gpt-4o-2024-11-20) – โมเดลหลักที่มีความสามารถสูงสุดในการประมวลผลข้อมูลทางคลินิก')
bullet('GPT-4o-mini – โมเดลที่มีต้นทุนต่ำกว่า เปรียบเทียบ cost-effectiveness')
bullet('ระบบ Baseline (GPT-4o ไม่มี Knowledge Base) – ใช้เป็น Baseline เพื่อวัดคุณค่าของการผสาน Knowledge Base')

h3('เกณฑ์การคัดเลือกโมเดล')
body('เกณฑ์การพิจารณาคัดเลือก LLM ที่จะนำมาใช้ในระบบ KBGI ประกอบด้วย')
bullet('ความสามารถในการประมวลผลข้อมูลทางการแพทย์และเภสัชกรรม (Medical NLP capability)')
bullet('ขนาด Context Window ที่เพียงพอสำหรับข้อมูลใบสั่งยาและ Knowledge Base (≥ 128,000 tokens)')
bullet('ความสามารถในการตอบสนองเป็นโครงสร้าง JSON (Structured Output) เพื่อรองรับ Multi-label Classification')
bullet('มี API ที่เสถียรและรองรับ Batch Processing สำหรับการวิเคราะห์ข้อมูลปริมาณมาก')
bullet('มีการรับรองด้านความปลอดภัยของข้อมูล (Data Privacy Compliance) ที่เหมาะสมสำหรับข้อมูลผู้ป่วย')

body(
    'การเปรียบเทียบระหว่างโมเดลจะใช้ชุดข้อมูลทดสอบเดียวกัน (Hold-out Test Set) '
    'ที่ผ่านการประเมินโดยเภสัชกรผู้เชี่ยวชาญ (Reference Standard) '
    'เพื่อให้ผลการเปรียบเทียบมีความยุติธรรมและน่าเชื่อถือ'
)

doc.add_paragraph()

# ── KBGI Development (instruction 2) ─────────────────────────────────────────
h2('การสร้างเครื่องมือที่ใช้ในการเก็บรวบรวมข้อมูล  [แก้ไขตามข้อ 2]')
note('ข้อ 2 – ขยายรายละเอียดกระบวนการสร้างระบบ Knowledge Base Generative AI')

h3('1. การพัฒนา Knowledge-Based Generative AI System (KBGI)')
body(
    'การพัฒนาระบบ KBGI ในงานวิจัยนี้มีวัตถุประสงค์เพื่อสร้างระบบปัญญาประดิษฐ์ที่ '
    'สามารถวิเคราะห์ข้อมูลใบสั่งยา เชื่อมโยงข้อมูลกับองค์ความรู้ทางคลินิก '
    'และจำแนกประเภทความผิดพลาดของใบสั่งยาในลักษณะ Multi-label Classification '
    'โดยอาศัยแนวคิดของ Knowledge-Based Generative AI ซึ่งผสานความสามารถของโมเดลภาษา '
    'ขนาดใหญ่ (Large Language Models) กับฐานความรู้ทางการแพทย์และเภสัชกรรม '
    'ที่ผ่านการตรวจสอบแล้ว'
)

h3('1.1 การรวบรวมฐานความรู้ (Knowledge Base Construction) – รายละเอียดขั้นตอน')
body('กระบวนการสร้างฐานความรู้ประกอบด้วย 4 ขั้นตอนหลัก ดังนี้')

bullet('ขั้นตอนที่ 1 – การรวบรวมแหล่งข้อมูล (Source Identification & Collection)')
body(
    'รวบรวมองค์ความรู้จากแหล่งข้อมูลทางการแพทย์ที่ได้รับการยอมรับในระดับสากล '
    'ได้แก่ UpToDate, Medscape, Thai Pharmacopoeia, สมาคมเภสัชกรรมโรงพยาบาล '
    '(ประเทศไทย) และวรรณกรรมวิชาการที่เกี่ยวข้อง ครอบคลุมหัวข้อ: '
    'คำจำกัดความ Prescription Error, แนวทางการใช้ยาอย่างเหมาะสม, '
    'ข้อมูลปฏิกิริยาระหว่างยา (Drug–Drug Interactions), '
    'หลักการปรับขนาดยาตามลักษณะผู้ป่วย, และข้อมูลข้อบ่งใช้/ข้อห้ามใช้ยา',
    indent=False
)

bullet('ขั้นตอนที่ 2 – การจัดโครงสร้างและจัดระเบียบข้อมูล (Knowledge Structuring)')
body(
    'แปลงข้อมูลจากแหล่งต่าง ๆ ให้อยู่ในรูปแบบโครงสร้าง (Structured Format) '
    'ที่ระบบ AI สามารถนำไปใช้ได้ โดยแบ่งเป็นหมวดหมู่ตามประเภทของ Prescription Error '
    'ได้แก่ Wrong Drug, Wrong Dose, Wrong Route, Wrong Frequency, '
    'Contraindication, Drug-Drug Interaction, และ Therapeutic Duplication '
    'จัดเก็บในรูปแบบ JSON-LD และ Vector Embedding เพื่อรองรับการค้นหาแบบ Semantic Search',
    indent=False
)

bullet('ขั้นตอนที่ 3 – การตรวจสอบความถูกต้อง (Knowledge Validation)')
body(
    'เภสัชกรผู้เชี่ยวชาญด้านเภสัชกรรมคลินิกอย่างน้อย 2 คน ทำการตรวจสอบความถูกต้อง '
    'ครบถ้วน และความเป็นปัจจุบันของข้อมูลในฐานความรู้ โดยใช้เกณฑ์ฉันทามติ '
    '(Consensus Criteria) กรณีที่มีความเห็นไม่ตรงกัน จะนำเข้าสู่กระบวนการพิจารณาโดย '
    'ผู้เชี่ยวชาญที่สาม (Third Expert Adjudication)',
    indent=False
)

bullet('ขั้นตอนที่ 4 – การอัปเดตและบำรุงรักษา (Knowledge Maintenance)')
body(
    'กำหนดแผนการทบทวนและอัปเดตฐานความรู้ทุก 3 เดือน หรือเมื่อมีการเปลี่ยนแปลง '
    'แนวทางการรักษาที่สำคัญ เพื่อให้ข้อมูลมีความเป็นปัจจุบันตลอดระยะเวลาการวิจัย',
    indent=False
)

doc.add_paragraph()

h3('1.2 การออกแบบระบบปัญญาประดิษฐ์ (AI System Design) – รายละเอียขั้นตอน')
body('ระบบ KBGI ถูกออกแบบโดยมีสถาปัตยกรรมหลัก 5 ส่วน ดังนี้')

bullet('ส่วนที่ 1 – Data Input Layer')
body(
    'รับข้อมูลใบสั่งยาจากระบบ HIS ในรูปแบบ HL7 FHIR หรือ CSV '
    'ผ่านกระบวนการ Anonymization เพื่อปกป้องข้อมูลส่วนบุคคลของผู้ป่วย '
    'ก่อนส่งเข้ากระบวนการวิเคราะห์',
    indent=False
)

bullet('ส่วนที่ 2 – Knowledge Retrieval Module (RAG)')
body(
    'ใช้เทคนิค Retrieval-Augmented Generation (RAG) ในการดึงข้อมูลที่เกี่ยวข้อง '
    'จากฐานความรู้ (Vector Database) โดยแปลงข้อมูลใบสั่งยาเป็น Query Vector '
    'และค้นหา Knowledge Chunks ที่มีความสัมพันธ์สูงสุด (Top-k Retrieval, k=5) '
    'เพื่อเสริมบริบทให้กับโมเดล AI',
    indent=False
)

bullet('ส่วนที่ 3 – Prompt Engineering Module')
body(
    'สร้าง Prompt ที่มีโครงสร้างชัดเจน (Structured Prompt) ประกอบด้วย: '
    '(1) System Prompt ที่กำหนดบทบาทของ AI เป็นเภสัชกรผู้เชี่ยวชาญ, '
    '(2) Knowledge Context จาก RAG Module, '
    '(3) Patient Data ข้อมูลผู้ป่วยและใบสั่งยา, '
    '(4) Task Instructions คำสั่งให้จำแนกประเภท Prescription Error แบบ Multi-label, '
    '(5) Output Format Specification กำหนดรูปแบบผลลัพธ์เป็น JSON',
    indent=False
)

bullet('ส่วนที่ 4 – LLM Inference Engine')
body(
    'ส่ง Prompt ที่สร้างแล้วไปยัง OpenAI API (Responses API) '
    'และรับผลลัพธ์ในรูปแบบ JSON ที่ระบุประเภทของ Prescription Error '
    'พร้อม Confidence Score และเหตุผลสนับสนุน (Reasoning)',
    indent=False
)

bullet('ส่วนที่ 5 – Output Processing & Storage')
body(
    'แปลงผลลัพธ์ JSON เป็น Multi-label Binary Matrix '
    'บันทึกผลลัพธ์ลงฐานข้อมูลพร้อม Log สำหรับการตรวจสอบย้อนหลัง '
    'และส่งออกในรูปแบบที่เหมาะสมสำหรับการวิเคราะห์ทางสถิติ',
    indent=False
)

doc.add_paragraph()

h3('1.3 การทดสอบระบบ (System Testing)')
body(
    'ก่อนนำระบบ KBGI ไปใช้ในการวิเคราะห์ข้อมูลจริง จะทำการทดสอบระบบใน 2 ระยะ:'
)
bullet('Pilot Testing: ทดสอบกับใบสั่งยาตัวอย่าง 30 ใบ ที่เภสัชกรประเมินแล้ว เพื่อตรวจสอบความถูกต้องของ System Pipeline')
bullet('Calibration Testing: ปรับ Prompt Template และ Knowledge Retrieval Parameters จนระบบมีประสิทธิภาพเบื้องต้นที่ยอมรับได้ (Precision ≥ 0.70)')

doc.add_paragraph()

# ── HIS Data Import (instruction 3) ───────────────────────────────────────────
h2('การนำเข้าข้อมูลจากระบบ HIS  [เพิ่มใหม่ตามข้อ 3]')
note('ข้อ 3 – อธิบายกระบวนการนำเข้าข้อมูลจาก HIS และขอบเขตที่เน้น Prescription Error')

body(
    'ข้อมูลที่ใช้ในการวิจัยนี้ดึงมาจากระบบสารสนเทศโรงพยาบาล '
    '(Hospital Information System; HIS) ของโรงพยาบาลกรุงเทพขอนแก่น '
    'ซึ่งเป็นระบบที่บันทึกข้อมูลการสั่งจ่ายยาทุกรายการในรูปแบบอิเล็กทรอนิกส์'
)

h3('ขั้นตอนการนำเข้าข้อมูล HIS')
bullet('ขั้นตอนที่ 1 – ฝ่ายเทคโนโลยีสารสนเทศของโรงพยาบาลดึงข้อมูลใบสั่งยาจากระบบ HIS ตามเกณฑ์การคัดเลือกที่กำหนด ในรูปแบบไฟล์ CSV หรือ HL7 FHIR')
bullet('ขั้นตอนที่ 2 – ทีมวิจัยทำการ De-identification โดยลบข้อมูลส่วนบุคคลที่ระบุตัวตนผู้ป่วยได้ เช่น ชื่อ นามสกุล HN หมายเลขคำสั่งยา และ Visit Number ก่อนส่งข้อมูลเข้าระบบ AI')
bullet('ขั้นตอนที่ 3 – ข้อมูลผ่าน Data Cleaning Pipeline ด้วยภาษา Python (ไลบรารี pandas) ตรวจสอบความครบถ้วน ลบข้อมูลซ้ำ และแปลงรูปแบบให้เหมาะสม')
bullet('ขั้นตอนที่ 4 – ส่งข้อมูลที่ผ่านการเตรียมแล้วเข้าสู่ระบบ KBGI เพื่อทำการวิเคราะห์และจำแนกประเภท Prescription Error')

body(
    'สำคัญ: ขอบเขตของข้อมูลที่วิเคราะห์ครอบคลุมเฉพาะ Prescription Error '
    '(ความผิดพลาดในขั้นตอนการสั่งยา) ซึ่งเป็นส่วนหนึ่งของ Medication Error ทั้งหมด '
    'ไม่รวมความผิดพลาดในขั้นตอนอื่น เช่น การเตรียมยา การจ่ายยา หรือการบริหารยา '
    'เนื่องจากข้อมูล HIS ที่เก็บได้ครอบคลุมเฉพาะข้อมูลการสั่งยาเป็นหลัก'
)

doc.add_paragraph()

# ── Data Analysis (instruction 7) ─────────────────────────────────────────────
h2('การวิเคราะห์ข้อมูลและสถิติ  [แก้ไขตามข้อ 7]')
note('ข้อ 7 – เพิ่มการคำนวณ Precision/F1 เปรียบเทียบกับเภสัชกร กรณีผลไม่เหมือนกัน และ baseline comparison')

h3('1. การประเมินประสิทธิภาพโมเดล AI (AI Performance Evaluation)')
body('ตัวชี้วัดการจำแนกประเภท (Classification Metrics) สำหรับ Multi-label Classification')
bullet('Precision (Micro/Macro Average) – วัดความถูกต้องของการระบุประเภท Prescription Error')
bullet('Recall/Sensitivity (Micro/Macro Average) – วัดความสามารถในการตรวจจับ Prescription Error ทั้งหมด')
bullet('F1-Score (Micro/Macro/Sample Average) – ค่ากลางระหว่าง Precision และ Recall')
bullet('Subset Accuracy – ความถูกต้องของการจำแนกครบทุก Label พร้อมกัน')
bullet('Hamming Loss – สัดส่วนของ Label ที่จำแนกผิดต่อจำนวน Label ทั้งหมด')

h3('2. การเปรียบเทียบกับการประเมินของเภสัชกร (Comparison with Pharmacist)')
body(
    'เภสัชกรผู้เชี่ยวชาญจำนวน 2 คน จะทำการประเมินใบสั่งยาในกลุ่มทดสอบ (Test Set) '
    'อย่างอิสระ (Blind to AI Results) เพื่อเป็น Reference Standard '
    'โดยผลการประเมินของเภสัชกรจะถูกนำมาเปรียบเทียบกับผล AI ในระดับ Label ดังนี้:'
)
bullet('Cohen\'s Kappa (κ) – วัดความสอดคล้องระหว่าง AI กับเภสัชกรในแต่ละ Label')
bullet('Inter-Rater Agreement (IRR) – ความสอดคล้องระหว่างเภสัชกรทั้งสองคนก่อนใช้เป็น Reference')
bullet('ROC-AUC per Label – ประเมินความสามารถในการแยกแยะ Positive/Negative ของแต่ละประเภท Error')

h3('3. การวิเคราะห์กรณีที่ผลไม่เหมือนกัน (Discordant Case Analysis)')
note('ข้อ 7 – การแปลผลที่ไม่เหมือนกันระหว่าง AI กับเภสัชกร')
body(
    'กรณีที่ระบบ AI จำแนกผลไม่ตรงกับเภสัชกร (Discordant Cases) '
    'จะถูกวิเคราะห์เพิ่มเติม ดังนี้:'
)
bullet('False Positive Analysis: กรณีที่ AI ระบุว่ามี Error แต่เภสัชกรไม่เห็นด้วย – วิเคราะห์สาเหตุ เช่น ข้อมูลใน Knowledge Base ไม่ครบ หรือ Context ไม่เพียงพอ')
bullet('False Negative Analysis: กรณีที่ AI ไม่ตรวจพบ Error ที่เภสัชกรระบุ – วิเคราะห์ประเภทของ Error ที่ระบบยังตรวจพบได้ยาก')
bullet('Adjudication: กรณีที่เภสัชกรทั้งสองคนไม่เห็นด้วยกัน จะนำเข้าสู่กระบวนการ Third Expert Adjudication ก่อนใช้เป็น Reference')
body(
    'ผลการวิเคราะห์ Discordant Cases จะถูกนำมาใช้ในการปรับปรุง Prompt Template '
    'และ Knowledge Base เพื่อพัฒนาความแม่นยำของระบบในระยะต่อไป'
)

h3('4. การเปรียบเทียบกับ Baseline (Without Knowledge Base)')
note('ข้อ 7 – เปรียบเทียบกับแบบไม่มี Knowledge Base AI')
body(
    'เพื่อประเมินคุณค่าของการผสาน Knowledge Base เข้ากับ Generative AI '
    'จะทำการเปรียบเทียบประสิทธิภาพระหว่าง 3 เงื่อนไข ดังนี้:'
)
bullet('เงื่อนไขที่ 1 – KBGI (With Knowledge Base): ระบบ GPT-4o + RAG Knowledge Base [ระบบหลักของการวิจัย]')
bullet('เงื่อนไขที่ 2 – LLM Only (Without Knowledge Base): ระบบ GPT-4o โดยไม่มี Knowledge Base [Baseline AI]')
bullet('เงื่อนไขที่ 3 – Pharmacist Assessment: การประเมินโดยเภสัชกรผู้เชี่ยวชาญ [Reference Standard]')
body(
    'การเปรียบเทียบจะใช้ Paired t-test หรือ McNemar\'s Test '
    'สำหรับเปรียบเทียบ Proportion ระหว่างเงื่อนไข '
    'ณ ระดับนัยสำคัญทางสถิติ p < 0.05 '
    'เพื่อตัดสินว่าการผสาน Knowledge Base ช่วยเพิ่มประสิทธิภาพอย่างมีนัยสำคัญหรือไม่'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# REFERENCES – with duplicate removed
# ═══════════════════════════════════════════════════════════════════════════════
h1('บรรณานุกรม  (แก้ไขตามข้อ 8 – ลบ Reference ซ้ำ)')
note('ข้อ 8 – Mesko & Topol (2023a) และ (2023b) มี DOI เดียวกัน (10.1038/s41746-023-00873-0) → ลบ 2023b ออก ใช้ 2023a เพียงรายการเดียว')
divider()

refs = [
    'Alsulami, Z., Conroy, S., & Choonara, I. (2013). Medication errors in the Middle East countries: A systematic review of the literature. European Journal of Clinical Pharmacology, 69(4), 995–1008. https://doi.org/10.1007/s00228-012-1435-y',
    'Amin, D., Garzón-Orjuela, N., Garcia Pereira, A., Parveen, S., Vornhagen, H., & Vellinga, A. (2023). Artificial Intelligence to Improve Antibiotic Prescribing: A Systematic Review. Antibiotics, 12(8), 1293. https://doi.org/10.3390/antibiotics12081293',
    'Balestra, M., Chen, J., Iturrate, E., Aphinyanaphongs, Y., & Nov, O. (2021). Predicting inpatient pharmacy order interventions using provider action data. JAMIA Open, 4(3), ooab083. https://doi.org/10.1093/jamiaopen/ooab083',
    'Bhatt, A., Roberts, R., Chen, X., Li, T., Connor, S., Hatim, Q., Mikailov, M., Tong, W., & Liu, Z. (2021). DICE: A Drug Indication Classification and Encyclopedia for AI-Based Indication Extraction. Frontiers in Artificial Intelligence, 4, 711467. https://doi.org/10.3389/frai.2021.711467',
    'Bishop, C. M. (2006). Pattern recognition and machine learning. Springer.',
    'Chelli, M., Descamps, J., Lavoue, V., Trojani, C., Azar, M., Deckert, M., Raynier, J.-L., Clowez, G., Boileau, P., & Ruetsch-Chelli, C. (2024). Hallucination Rates and Reference Accuracy of ChatGPT and Bard for Systematic Reviews: Comparative Analysis. Journal of Medical Internet Research, 26, e53164. https://doi.org/10.2196/53164',
    'Corny, J., Rajkumar, A., Martin, O., Dode, X., Lajonchere, J.-P., Billuart, O., Bezie, Y., & Buronfosse, A. (2020). A machine learning–based clinical decision support system to identify prescriptions with a high risk of medication error. Journal of the American Medical Informatics Association, 27(11), 1688–1694. https://doi.org/10.1093/jamia/ocaa154',
    'Coulet, A., Shah, N. H., Wack, M., Chawki, M. B., Jay, N., & Dumontier, M. (2018). Predicting the need for a reduced drug dose, at first prescription. Scientific Reports, 8(1), 15558. https://doi.org/10.1038/s41598-018-33980-0',
    'Crossnohere, N. L., Elsaid, M., Paskett, J., Bose-Brill, S., & Bridges, J. F. P. (2022). Guidelines for Artificial Intelligence in Medicine: Literature Review and Content Analysis of Frameworks. Journal of Medical Internet Research, 24(8), e36823. https://doi.org/10.2196/36823',
    'Gatto, E. C., Ferrandin, M., & Cerri, R. (2025). Multi-label classification with label clusters. Knowledge and Information Systems, 67(2), 1741–1785. https://doi.org/10.1007/s10115-024-02270-9',
    'Gonzaga De Andrade Santos, T. N., Mendonça Da Cruz Macieira, G., Cardoso Sodre Alves, B. M., Onozato, T., Cunha Cardoso, G., Ferreira Nascimento, M. T., Saquete Martins-Filho, P. R., Pereira De Lyra, D., & Oliveira Filho, A. D. D. (2020). Prevalence of clinically manifested drug interactions in hospitalized patients: A systematic review and meta-analysis. PLOS ONE, 15(7), e0235353. https://doi.org/10.1371/journal.pone.0235353',
    'Härkänen, M., Haatainen, K., Vehviläinen-Julkunen, K., & Miettinen, M. (2021). Artificial Intelligence for Identifying the Prevention of Medication Incidents Causing Serious or Moderate Harm: An Analysis Using Incident Reporters\' Views. International Journal of Environmental Research and Public Health, 18(17), 9206. https://doi.org/10.3390/ijerph18179206',
    'Heaton, J. (2018). Ian Goodfellow, Yoshua Bengio, and Aaron Courville: Deep learning. Genetic Programming and Evolvable Machines, 19(1–2), 305–307. https://doi.org/10.1007/s10710-017-9314-z',
    'Iancu, A., Leb, I., Prokosch, H.-U., & Rödle, W. (2023). Machine learning in medication prescription: A systematic review. International Journal of Medical Informatics, 180, 105241. https://doi.org/10.1016/j.ijmedinf.2023.105241',
    'Klumpp, M., et al. (2021). Artificial Intelligence for Hospital Health Care: Application Cases and Answers to Challenges in European Hospitals. Healthcare, 9(8), 961. https://doi.org/10.3390/healthcare9080961',
    'Kuperman, G. J., Bobb, A., Payne, T. H., Avery, A. J., Gandhi, T. K., Burns, G., Classen, D. C., & Bates, D. W. (2007). Medication-related Clinical Decision Support in Computerized Provider Order Entry Systems: A Review. Journal of the American Medical Informatics Association, 14(1), 29–40. https://doi.org/10.1197/jamia.M2170',
    'Lan, Y., He, S., Liu, K., Zeng, X., Liu, S., & Zhao, J. (2021). Path-based knowledge reasoning with textual semantic information for medical knowledge graph completion. BMC Medical Informatics and Decision Making, 21(S9), 335. https://doi.org/10.1186/s12911-021-01622-7',
    'Liu, S., Kawamoto, K., Del Fiol, G., Weir, C., Malone, D. C., Reese, T. J., Morgan, K., ElHalta, D., & Abdelrahman, S. (2022). The potential for leveraging machine learning to filter medication alerts. Journal of the American Medical Informatics Association: JAMIA, 29(5), 891–899. https://doi.org/10.1093/jamia/ocab292',
    'Luo, H., Yin, W., Wang, J., Zhang, G., Liang, W., Luo, J., & Yan, C. (2024). Drug-drug interactions prediction based on deep learning and knowledge graph: A review. iScience, 27(3), 109148. https://doi.org/10.1016/j.isci.2024.109148',
    'Lyu, K., Tian, Y., Shang, Y., Zhou, T., Yang, Z., Liu, Q., Yao, X., Zhang, P., Chen, J., & Li, J. (2023). Causal knowledge graph construction and evaluation for clinical decision support of diabetic nephropathy. Journal of Biomedical Informatics, 139, 104298. https://doi.org/10.1016/j.jbi.2023.104298',
    'Mahomedradja, R. F., Schinkel, M., Sigaloff, K. C. E., Reumerman, M. O., Otten, R. H. J., Tichelaar, J., & Van Agtmael, M. A. (2023). Factors influencing in-hospital prescribing errors: A systematic review. British Journal of Clinical Pharmacology, 89(6), 1724–1735. https://doi.org/10.1111/bcp.15694',
    'Mesko, B., & Topol, E. J. (2023). The imperative for regulatory oversight of large language models (or generative AI) in healthcare. Npj Digital Medicine, 6(1), 120. https://doi.org/10.1038/s41746-023-00873-0',
    'Natsiavas, P., Malousi, A., Bousquet, C., Jaulent, M.-C., & Koutkias, V. (2019). Computational Advances in Drug Safety: Systematic and Mapping Review of Knowledge Engineering Based Approaches. Frontiers in Pharmacology, 10, 415. https://doi.org/10.3389/fphar.2019.00415',
    'Noorbakhsh-Sabet, N., Zand, R., Zhang, Y., & Abedi, V. (2019). Artificial Intelligence Transforms the Future of Health Care. The American Journal of Medicine, 132(7), 795–801. https://doi.org/10.1016/j.amjmed.2019.01.017',
    'Patel, J., Ladani, A., Sambamoorthi, N., LeMasters, T., Dwibedi, N., & Sambamoorthi, U. (2020). A Machine Learning Approach to Identify Predictors of Potentially Inappropriate Non-Steroidal Anti-Inflammatory Drugs (NSAIDs) Use in Older Adults with Osteoarthritis. International Journal of Environmental Research and Public Health, 18(1), 155. https://doi.org/10.3390/ijerph18010155',
    'Poly, T. N., Islam, M. M., Muhtar, M. S., Yang, H.-C., Nguyen, P. A. A., & Li, Y.-C. J. (2020). Machine Learning Approach to Reduce Alert Fatigue Using a Disease Medication-Related Clinical Decision Support System. JMIR Medical Informatics, 8(11), e19489. https://doi.org/10.2196/19489',
    'Rajkomar, A., Dean, J., & Kohane, I. (2019). Machine Learning in Medicine. New England Journal of Medicine, 380(14), 1347–1358. https://doi.org/10.1056/NEJMra1814259',
    'Ranchon, F., Chanoine, S., Lambert-Lacroix, S., Bosson, J.-L., Moreau-Gaudry, A., & Bedouch, P. (2023). Development of artificial intelligence powered apps and tools for clinical pharmacy services: A systematic review. International Journal of Medical Informatics, 172, 104983. https://doi.org/10.1016/j.ijmedinf.2022.104983',
    'Singh, B., Olds, T., Brinsley, J., Dumuid, D., Virgara, R., Matricciani, L., Watson, A., Szeto, K., Eglitis, E., Miatke, A., Simpson, C. E. M., Vandelanotte, C., & Maher, C. (2023). Systematic review and meta-analysis of the effectiveness of chatbots on lifestyle behaviours. Npj Digital Medicine, 6(1), 118. https://doi.org/10.1038/s41746-023-00856-1',
    'Soroush, A., Glicksberg, B. S., Zimlichman, E., Barash, Y., Freeman, R., Charney, A. W., Nadkarni, G. N., & Klang, E. (2024). Large Language Models Are Poor Medical Coders—Benchmarking of Medical Code Querying. NEJM AI, 1(5). https://doi.org/10.1056/Aldbp2300040',
    'Sutton, R. T., Pincock, D., Baumgart, D. C., Sadowski, D. C., Fedorak, R. N., & Kroeker, K. I. (2020). An overview of clinical decision support systems: Benefits, risks, and strategies for success. Npj Digital Medicine, 3(1), 17. https://doi.org/10.1038/s41746-020-0221-y',
    'The Lancet Digital Health. (2023). ChatGPT: Friend or foe? The Lancet Digital Health, 5(3), e102. https://doi.org/10.1016/S2589-7500(23)00023-7',
    'Topol, E. J. (2019). High-performance medicine: The convergence of human and artificial intelligence. Nature Medicine, 25(1), 44–56. https://doi.org/10.1038/s41591-018-0300-7',
    'Tsoumakas, G., & Katakis, I. (n.d.). Multi-Label Classification: An Overview.',
    'Zhang, M.-L., & Zhou, Z.-H. (2014). A Review on Multi-Label Learning Algorithms. IEEE Transactions on Knowledge and Data Engineering, 26(8), 1819–1837. https://doi.org/10.1109/TKDE.2013.39',
]

for ref in refs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.first_line_indent = Cm(-1.25)
    run = p.add_run(ref)
    run.font.size = Pt(14)

doc.add_paragraph()
note('NOTE: Mesko & Topol 2023b (เดิม) ถูกลบออกเนื่องจากเป็น Reference เดียวกับ 2023a (DOI: 10.1038/s41746-023-00873-0 ซ้ำกัน) — บทความที่อ้างในเนื้อหาใช้เพียง (Mesko & Topol, 2023a) หรือ (Mesko & Topol, 2023)')

# ── Save ───────────────────────────────────────────────────────────────────────
out = r'F:\Shared_folders\KBAI-thesis\docs\thesis\Thesis_Proposal_Revised_660820006.docx'
doc.save(out)
print(f'Saved: {out}')
